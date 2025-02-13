from docx.shared import Cm
import docx
import csv
import whisper
import re
import os
import sys

# dir_path = './'
extensions = [".mp4", ".mov", "mkv", ".webm",
              ".mp3", ".wav", ".flac", ".aac", ".ogg", ".m4a"]

files = os.listdir()
file_name = ''
file_extension = ''
for extension in extensions:
    for name in files:
        if name.endswith(extension):
            file_name = name
            file_extension = extension
            break
        if file_name != '':
            break
if file_name == '':
    print('対応しているファイル形式が見つかりませんでした')
    sys.exit(0)


filemeta_name = file_name.replace(file_extension, '')


date_match = re.match(r"(\d{4})(\d{2})(\d{2})", filemeta_name)
if not date_match:
    print("ファイル名が不適切です。")
    print("YYYYMMDD_HHMM　の形式にしてください。")
    sys.exit(0)


# modelのロード（large-v2を選択する場合）
model = whisper.load_model("large-v2")
# 文字起こし
result = model.transcribe(file_name, verbose=True)


with open("台本_" + filemeta_name + ".csv", 'w', newline='') as csvfile:

    writer = csv.writer(csvfile)

    writer.writerows([['start_time', 'end_time', 'output']])
    for segment in result['segments']:
        ts = int(segment['start'])
        te = int(segment['end'])
        start = f"{ts%60}:{ts//60}"
        end = f"{te%60}:{ts//60}"
        text = segment['text']
        writer.writerows([[start, end, text]])


# Word文書の新規作成
doc = docx.Document()

doc.add_heading(filemeta_name[14:])
doc.add_paragraph(f"{date_match[1]}年{date_match[2]}月{date_match[3]}日放送")

# CSVファイルを読み込み
with open("台本_" + filemeta_name + ".csv", 'w', newline='') as csvfile:

    writer = csv.writer(csvfile)

    writer.writerows([['start_time', 'end_time', 'output']])
    for segment in result['segments']:
        ts = int(segment['start'])
        te = int(segment['end'])
        start = f"{ts//60:02}:{ts%60:02}"
        end = f"{te//60:02}:{te%60:02}"
        text = segment['text']
        writer.writerows([[start, end, text]])


# Word文書の新規作成
doc = docx.Document()

doc.add_heading(filemeta_name[14:])
doc.add_paragraph(f"{date_match[1]}年{date_match[2]}月{date_match[3]}日放送")

# CSVファイルを読み込み
with open("台本_" + filemeta_name + ".csv", encoding='utf-8') as f:
    reader = csv.reader(f)
    data = [row for row in reader]

# 表の作成
table = doc.add_table(rows=len(data), cols=len(data[0]))
table.style = 'Table Grid'  # 罫線を追加

# データを表に挿入
for r, row in enumerate(data):
    for c, cell in enumerate(row):
        table.cell(r, c).text = cell


def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width


table.autofit = False
table.allow_autofit = False

set_column_width(table.columns[0], Cm(2.2))
set_column_width(table.columns[1], Cm(2.2))
set_column_width(table.columns[2], Cm(11.2))

# Word文書の保存
doc.save("台本_" + filemeta_name + ".docx")


file_path = file_name  # 削除したいファイルのパス
if os.path.exists(file_path):
    os.remove(file_path)
    print("次の動画をアップロードできます。")
else:
    print("ファイルが見つかりません。")
